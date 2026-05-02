"""DISE (ASU 2024-03) exports — CSV / JSON / DOCX.

Builds the natural-expense-by-caption pivot table from approved DISE
mappings in Supabase, then renders it in three formats. The DOCX
includes the four required ASU 2024-03 narrative sections (methodology,
selling expenses definition, inventory expensed when sold, election
change).
"""

from __future__ import annotations

import csv
import io
import json
import logging
from typing import Any

log = logging.getLogger("exports.dise")

DISE_CATEGORIES = [
    "Purchases of inventory",
    "Employee compensation",
    "Depreciation",
    "Intangible asset amortization",
    "Other expenses",
]
DISE_CAPTIONS = ["COGS", "SG&A", "R&D", "Other income/expense"]
DISE_CITATIONS = {
    "Purchases of inventory":        "ASC 220-40-50-6(b)",
    "Employee compensation":         "ASC 220-40-50-6(a)",
    "Depreciation":                  "ASC 220-40-50-6(c)",
    "Intangible asset amortization": "ASC 220-40-50-6(d)",
    "Other expenses":                "ASC 220-40-50-6(e)",
}


def build_pivot(rows: list[dict]) -> tuple[dict, dict, float]:
    """Return (pivot[cat][cap], col_totals[cap], grand_total)."""
    pivot: dict[str, dict[str, float]] = {c: {cap: 0.0 for cap in DISE_CAPTIONS} for c in DISE_CATEGORIES}
    for r in rows:
        cat = r.get("dise_category")
        cap = r.get("expense_caption")
        if cat in pivot and cap in pivot[cat]:
            pivot[cat][cap] += float(r.get("posting_amount", 0))
    col_totals = {cap: sum(pivot[c][cap] for c in DISE_CATEGORIES) for cap in DISE_CAPTIONS}
    grand_total = sum(col_totals.values())
    return pivot, col_totals, grand_total


def to_csv(rows: list[dict]) -> str:
    """Per-account approved-mappings CSV — what an external auditor will inspect."""
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow([
        "gl_account", "description", "posting_amount", "dise_category",
        "expense_caption", "asc_citation", "override_reason",
        "reviewer", "reviewed_at",
    ])
    for r in rows:
        writer.writerow([
            r.get("gl_account", ""),
            (r.get("description") or "").replace("\n", " "),
            float(r.get("posting_amount", 0)),
            r.get("dise_category", ""),
            r.get("expense_caption", ""),
            r.get("asc_citation", ""),
            (r.get("override_reason") or ""),
            r.get("reviewer", ""),
            r.get("reviewed_at", ""),
        ])
    return buf.getvalue()


def to_json(rows: list[dict], fiscal_year: str) -> str:
    pivot, col_totals, grand_total = build_pivot(rows)
    return json.dumps({
        "fiscal_year": fiscal_year,
        "approved_count": len(rows),
        "rows": rows,
        "pivot": pivot,
        "column_totals": col_totals,
        "grand_total": grand_total,
    }, indent=2, default=str)


def to_docx(rows: list[dict], fiscal_year: str) -> bytes:
    """Build the Workiva-ready DISE footnote DOCX. Amounts in thousands."""
    from docx import Document  # type: ignore
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.shared import Pt

    pivot, col_totals, grand_total = build_pivot(rows)

    def fmt_thousands(n: float) -> str:
        v = round(n / 1000)
        return f"{v:,}" if v else "—"

    doc = Document()
    doc.add_heading(
        f"Disaggregation of Income Statement Expenses (DISE) — FY{fiscal_year}",
        level=1,
    )
    doc.add_paragraph(
        f"Tabular disclosure required by ASU 2024-03 (ASC 220-40). "
        f"Amounts in thousands. Built from {len(rows)} controller-approved "
        "GL classifications."
    )

    # Pivot table: rows × captions + totals.
    headers = ["Natural expense category", "ASC citation"] + DISE_CAPTIONS + ["Total"]
    table = doc.add_table(rows=1 + len(DISE_CATEGORIES) + 1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for col, h in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, cat in enumerate(DISE_CATEGORIES, start=1):
        row_total = sum(pivot[cat][cap] for cap in DISE_CAPTIONS)
        cells = table.rows[i].cells
        cells[0].text = cat
        cells[1].text = DISE_CITATIONS[cat]
        for j, cap in enumerate(DISE_CAPTIONS, start=2):
            cells[j].text = fmt_thousands(pivot[cat][cap])
        cells[-1].text = fmt_thousands(row_total)

    total_row = table.rows[-1]
    total_row.cells[0].text = "Total"
    for j, cap in enumerate(DISE_CAPTIONS, start=2):
        total_row.cells[j].text = fmt_thousands(col_totals[cap])
    total_row.cells[-1].text = fmt_thousands(grand_total)
    for c in total_row.cells:
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True

    # Required ASU 2024-03 narrative sections.
    doc.add_heading("Methodology", level=2)
    doc.add_paragraph(
        "The Company classified its income statement expenses into the natural "
        "expense categories prescribed by ASC 220-40 using a controller-reviewed "
        "AI classification of source GL accounts. Each classification is supported "
        "by the underlying GL detail and was approved by the Controller prior to "
        "inclusion in this disclosure. Reasonable estimates and methods that "
        "approximate the prescribed categories were used where source detail was "
        "not directly available, as permitted by ASC 220-40."
    )
    doc.add_heading("Selling expenses definition", level=2)
    doc.add_paragraph(
        "For purposes of this disclosure, selling expenses include direct sales "
        "force compensation and benefits, advertising and marketing programs, "
        "third-party sales commissions, sales support technology, and travel "
        "incurred in pursuit of new and existing customer relationships. The "
        "Company applied this definition consistently across the periods presented."
    )
    doc.add_heading("Inventory expensed when sold", level=2)
    doc.add_paragraph(
        "Amounts of inventory expensed when sold are included in the "
        "\"Purchases of inventory\" row when those amounts are recognized in cost "
        "of revenues. Capitalized inventory is excluded from this disclosure "
        "until expensed."
    )
    doc.add_heading("Election change", level=2)
    doc.add_paragraph(
        "No election changes were made during the period presented. If an election "
        "is changed in a future period, the reason for the change will be disclosed "
        "and prior periods recast for comparative purposes, unless impracticable."
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
