"""Income Tax (ASU 2023-09) exports — CSV / JSON / DOCX.

Builds the simplified ETR reconciliation + pre-tax split tables from
approved tax mappings, plus a CSV of the per-account approvals for the
auditor evidence package.
"""

from __future__ import annotations

import csv
import io
import json
import logging
from typing import Any

log = logging.getLogger("exports.tax")

STATUTORY_RATE = 0.21


def _aggregate(rows: list[dict]) -> dict[str, float]:
    totals: dict[str, float] = {
        "current_federal": 0, "current_state": 0, "current_foreign": 0,
        "deferred_federal": 0, "deferred_state": 0, "deferred_foreign": 0,
        "deferred_tax_asset": 0, "deferred_tax_liab": 0,
        "pretax_domestic": 0, "pretax_foreign": 0,
    }
    for r in rows:
        cat = r.get("tax_category")
        if cat in totals:
            totals[cat] += float(r.get("posting_amount", 0))
    return totals


def to_csv(rows: list[dict]) -> str:
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow([
        "gl_account", "description", "posting_amount", "tax_category",
        "tax_category_label", "asc_citation", "override_reason",
        "reviewer", "reviewed_at",
    ])
    for r in rows:
        w.writerow([
            r.get("gl_account", ""),
            (r.get("description") or "").replace("\n", " "),
            float(r.get("posting_amount", 0)),
            r.get("tax_category", ""),
            r.get("tax_category_label", ""),
            r.get("asc_citation", ""),
            r.get("override_reason") or "",
            r.get("reviewer", ""),
            r.get("reviewed_at", ""),
        ])
    return buf.getvalue()


def to_json(rows: list[dict], fiscal_year: str) -> str:
    totals = _aggregate(rows)
    pretax = totals["pretax_domestic"] + totals["pretax_foreign"]
    current_total = totals["current_federal"] + totals["current_state"] + totals["current_foreign"]
    deferred_total = totals["deferred_federal"] + totals["deferred_state"] + totals["deferred_foreign"]
    provision = current_total + deferred_total
    return json.dumps({
        "fiscal_year": fiscal_year,
        "approved_count": len(rows),
        "totals_by_category": totals,
        "current_provision_total": current_total,
        "deferred_provision_total": deferred_total,
        "total_provision": provision,
        "pretax_total": pretax,
        "effective_rate": (provision / pretax) if pretax else 0,
        "statutory_rate": STATUTORY_RATE,
    }, indent=2, default=str)


def to_docx(rows: list[dict], fiscal_year: str) -> bytes:
    from docx import Document  # type: ignore
    from docx.shared import Pt

    totals = _aggregate(rows)
    pretax = totals["pretax_domestic"] + totals["pretax_foreign"]
    current_total = totals["current_federal"] + totals["current_state"] + totals["current_foreign"]
    deferred_total = totals["deferred_federal"] + totals["deferred_state"] + totals["deferred_foreign"]
    provision = current_total + deferred_total
    effective = (provision / pretax) if pretax else 0
    state_local = totals["current_state"] + totals["deferred_state"]
    foreign_total = totals["current_foreign"] + totals["deferred_foreign"]
    foreign_diff = foreign_total - (totals["pretax_foreign"] * STATUTORY_RATE)
    deferred_federal = totals["deferred_federal"]

    fm = lambda n: f"${n / 1_000_000:.1f}M"
    fp = lambda r: f"{r * 100:.2f}%"

    doc = Document()
    doc.add_heading(f"Income Taxes (ASU 2023-09) — FY{fiscal_year}", level=1)
    doc.add_paragraph(
        f"ETR reconciliation, pre-tax income split, and supporting detail built "
        f"from {len(rows)} controller-approved tax GL classifications. Effective "
        f"tax rate: {fp(effective)}; statutory rate: {STATUTORY_RATE * 100:.0f}%."
    )

    doc.add_heading("Table A — Effective tax rate reconciliation", level=2)
    table_a_rows = [
        ("Item", "Amount", "% of pretax", True),
        (f"US federal statutory @{int(STATUTORY_RATE * 100)}%", fm(pretax * STATUTORY_RATE),
            fp(STATUTORY_RATE), False),
        ("State and local, net of federal benefit", fm(state_local),
            fp(state_local / pretax) if pretax else "—", False),
        ("Foreign rate differential", fm(foreign_diff),
            fp(foreign_diff / pretax) if pretax else "—", False),
        ("Deferred tax expense — federal", fm(deferred_federal),
            fp(deferred_federal / pretax) if pretax else "—", False),
        ("Total income tax expense", fm(provision), fp(effective), False),
    ]
    t = doc.add_table(rows=len(table_a_rows), cols=3)
    t.style = "Light Grid Accent 1"
    for ri, (item, amt, pct, hdr) in enumerate(table_a_rows):
        cells = t.rows[ri].cells
        cells[0].text = item
        cells[1].text = amt
        cells[2].text = pct
        if hdr:
            for c in cells:
                for p in c.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(10)

    doc.add_heading("Table C — Pre-tax income split", level=2)
    doc.add_paragraph(f"Domestic operations: {fm(totals['pretax_domestic'])}.")
    doc.add_paragraph(f"Foreign operations: {fm(totals['pretax_foreign'])}.")
    doc.add_paragraph(f"Total: {fm(pretax)}.")

    doc.add_heading("Table B — Cash income taxes paid", level=2)
    doc.add_paragraph(
        "Cash taxes paid disaggregation requires the prior-period cash-tax "
        "facts table; populate `tax_provision_data.cash_taxes_paid` to render "
        "this table."
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
